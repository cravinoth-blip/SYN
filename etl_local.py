"""
Local ETL Pipeline - Converts documents to embeddings and uploads to Snowflake.
Adapted from ETL_Code_Dec.ipynb (Google Colab version) for local execution.
"""

import os, json, re, csv, logging, time
from pathlib import Path
from dotenv import load_dotenv
from openai import OpenAI
from langchain_text_splitters import RecursiveCharacterTextSplitter
import pymupdf

# Optional heavy dependencies (only needed for Snowflake ETL upload, not for app.py)
try:
    from cryptography.hazmat.backends import default_backend
    from cryptography.hazmat.primitives.serialization import load_pem_private_key, Encoding, PrivateFormat, NoEncryption
    CRYPTO_AVAILABLE = True
except ImportError:
    CRYPTO_AVAILABLE = False

try:
    import pandas as pd
    PANDAS_AVAILABLE = True
except ImportError:
    PANDAS_AVAILABLE = False

try:
    from tqdm import tqdm
    TQDM_AVAILABLE = True
except ImportError:
    TQDM_AVAILABLE = False
    def tqdm(x, **kw):
        return x

try:
    import snowflake.connector as sf
    from snowflake.connector.pandas_tools import write_pandas
    SNOWFLAKE_AVAILABLE = True
except ImportError:
    SNOWFLAKE_AVAILABLE = False

# Optional imports
try:
    from docx import Document
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("python-docx not available. DOCX files will be skipped.")

try:
    import openpyxl
    EXCEL_AVAILABLE = True
except ImportError:
    EXCEL_AVAILABLE = False
    print("openpyxl not available. Excel files will be skipped.")

# ============================================================
# CONFIGURATION
# ============================================================
# Load env vars manually to skip the malformed SNOWFLAKE_PRIVATE_KEY line
_env_path = Path(__file__).parent / ".env"
if _env_path.exists():
    with open(_env_path, 'rb') as _f:
        for _line in _f.read().split(b'\n'):
            _line = _line.rstrip(b'\r').decode('utf-8', errors='ignore').strip()
            if '=' in _line and not _line.startswith('SNOWFLAKE_PRIVATE_KEY"') and not _line.startswith('#'):
                _k, _, _v = _line.partition('=')
                _k = _k.strip()
                _v = _v.strip().strip('"')
                if _k and _k == _k.upper() and not os.environ.get(_k):
                    os.environ[_k] = _v

# Paths - local folders
INPUT_FOLDER = Path(__file__).parent / "documents"
OUTPUT_FOLDER = Path(__file__).parent / "output"
CSV_FILENAME = "embeddings.csv"
OUTPUT_FOLDER.mkdir(exist_ok=True)

# Snowflake
TABLE_NAME = "RAG_DOCUMENTS"
UPLOAD_MODE = "replace"       # "replace", "append", or "upsert"
RECREATE_TABLE = True
UPLOAD_TO_SNOWFLAKE = True
GENERATE_NEW_CSV = False

# Chunking
CHUNK_SIZE = 1500
CHUNK_OVERLAP = 100
EMBEDDING_MODEL = "text-embedding-3-small"

# Snowflake connection config (private key auth)
def _load_private_key():
    pem_path = Path(__file__).parent / "private_key.pem"
    passphrase = os.getenv("SNOWFLAKE_PRIVATE_KEY_PASSPHRASE", "").encode()
    try:
        with open(pem_path, "rb") as f:
            pem_data = f.read()
        key = load_pem_private_key(pem_data, password=passphrase or None, backend=default_backend())
        return key.private_bytes(encoding=Encoding.DER, format=PrivateFormat.PKCS8, encryption_algorithm=NoEncryption())
    except Exception as e:
        print(f"Warning: Could not load private key: {e}")
        return None

SNOW_CONFIG = {
    'account': os.getenv("SNOWFLAKE_ACCOUNT"),
    'user': os.getenv("SNOWFLAKE_USER"),
    'private_key': _load_private_key(),
    'warehouse': os.getenv("SNOWFLAKE_WAREHOUSE"),
    'database': os.getenv("SNOWFLAKE_DATABASE"),
    'schema': os.getenv("SNOWFLAKE_SCHEMA"),
}

OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")

# ============================================================
# CONNECTION MANAGEMENT
# ============================================================
def get_snowflake_connection():
    try:
        ctx = sf.connect(**SNOW_CONFIG)
        ctx.cursor().execute("USE WAREHOUSE WH_COMMUNICATIONS__EU__DER")
        print("Connected to Snowflake successfully")
        return ctx
    except sf.errors.Error as e:
        print(f"Failed to connect to Snowflake: {e}")
        return None

def test_snowflake_connection():
    ctx = get_snowflake_connection()
    if ctx:
        try:
            cur = ctx.cursor()
            cur.execute("SELECT current_version();")
            version = cur.fetchone()[0]
            print(f"Snowflake version: {version}")
            return True
        except Exception as e:
            print(f"Connection test failed: {e}")
            return False
        finally:
            cur.close()
            ctx.close()
    return False

def get_openai_client():
    try:
        client = OpenAI(api_key=OPENAI_API_KEY)
        client.models.list()
        print("OpenAI client initialized successfully")
        return client
    except Exception as e:
        print(f"Failed to initialize OpenAI client: {e}")
        return None

# ============================================================
# FILE EXTRACTION FUNCTIONS
# ============================================================
def extract_content_from_pdf(pdf_path):
    print(f"Extracting: {pdf_path.name}")
    text_data = []
    document_text = ""
    try:
        with pymupdf.open(pdf_path) as pdf_doc:
            for page_number in range(pdf_doc.page_count):
                page = pdf_doc.load_page(page_number)
                page_text = page.get_text("text") or ""
                document_text += page_text + "\n"
                lines = page_text.splitlines()
                table_text = []
                non_table_text = []
                for line in lines:
                    if re.search(r"\s{4,}", line):
                        table_text.append(line)
                    else:
                        non_table_text.append(line)
                if table_text:
                    text_data.append({'text': "\n".join(table_text), 'page_number': page_number + 1, 'is_table': True})
                if non_table_text:
                    text_data.append({'text': "\n".join(non_table_text), 'page_number': page_number + 1, 'is_table': False})
        return {'filename': pdf_path.name, 'title': '', 'text_data': text_data, 'full_text': document_text}
    except Exception as e:
        print(f"Failed to extract from {pdf_path}: {e}")
        return {}

def extract_content_from_docx(docx_path):
    if not DOCX_AVAILABLE:
        return {}
    print(f"Extracting: {docx_path.name}")
    try:
        doc = docx.Document(docx_path)
        text_data = []
        full_text = ""
        for para in doc.paragraphs:
            if para.text.strip():
                full_text += para.text + "\n"
        if full_text:
            text_data.append({'text': full_text, 'page_number': 1, 'is_table': False})
        for table in doc.tables:
            table_text = "\n".join([" | ".join(cell.text.strip() for cell in row.cells) for row in table.rows])
            text_data.append({'text': table_text, 'page_number': 1, 'is_table': True})
        title = docx_path.stem
        return {'filename': docx_path.name, 'title': title, 'text_data': text_data, 'full_text': full_text}
    except Exception as e:
        print(f"Error extracting from {docx_path}: {e}")
        return {}

def extract_content_from_txt(txt_path):
    print(f"Extracting: {txt_path.name}")
    try:
        with open(txt_path, 'r', encoding='utf-8') as f:
            text = f.read()
        text = re.sub(r'\s+', ' ', text)
        return {
            'filename': txt_path.name, 'title': txt_path.stem,
            'text_data': [{'text': text, 'page_number': 1, 'is_table': False}],
            'full_text': text
        }
    except Exception as e:
        print(f"Error extracting from {txt_path}: {e}")
        return {}

def extract_content_from_json(json_path):
    print(f"Extracting: {json_path.name}")
    try:
        with open(json_path, 'r') as f:
            data = json.load(f)
        full_text = json.dumps(data, indent=2)
        return {
            'filename': json_path.name, 'title': json_path.stem,
            'text_data': [{'text': full_text, 'page_number': 1, 'is_table': False}],
            'full_text': full_text
        }
    except Exception as e:
        print(f"Error extracting from {json_path}: {e}")
        return {}

def extract_content_from_excel(excel_path):
    if not EXCEL_AVAILABLE:
        return {}
    print(f"Extracting: {excel_path.name}")
    try:
        workbook = openpyxl.load_workbook(excel_path, data_only=True)
        text_data = []
        all_text = ""
        for sheet_name in workbook.sheetnames:
            sheet = workbook[sheet_name]
            if sheet.max_row == 0:
                continue
            merged_cells_map = {}
            for merged_range in sheet.merged_cells.ranges:
                min_col, min_row, max_col, max_row = merged_range.bounds
                master_value = sheet.cell(min_row, min_col).value
                for row in range(min_row, max_row + 1):
                    for col in range(min_col, max_col + 1):
                        merged_cells_map[(row, col)] = master_value
            rows_data = []
            for row_idx, row in enumerate(sheet.iter_rows(min_row=1, max_row=sheet.max_row), start=1):
                row_values = []
                for col_idx, cell in enumerate(row, start=1):
                    value = merged_cells_map.get((row_idx, col_idx), cell.value)
                    row_values.append(value)
                rows_data.append(row_values)
            if not rows_data:
                continue
            df = pd.DataFrame(rows_data)
            header_row_idx = 0
            for idx, row in df.iterrows():
                if row.notna().any():
                    header_row_idx = idx
                    break
            if header_row_idx < len(df):
                headers = df.iloc[header_row_idx].fillna('').astype(str).tolist()
                headers = [str(h).strip() if str(h).strip() else f"Column_{i}" for i, h in enumerate(headers)]
                df = df.iloc[header_row_idx + 1:].copy()
                df.columns = headers
                df.reset_index(drop=True, inplace=True)
            df = df.dropna(how='all')
            df = df.loc[:, df.notna().any()]
            if df.empty:
                continue
            for original_idx, row in df.iterrows():
                row_text_parts = []
                for col in df.columns:
                    try:
                        cell_value = row[col]
                        if isinstance(cell_value, pd.Series):
                            cell_value = cell_value.iloc[0] if len(cell_value) > 0 else None
                        if pd.notna(cell_value):
                            cell_str = str(cell_value).strip()
                            if cell_str and cell_str.lower() not in ['nan', 'none', '']:
                                row_text_parts.append(f"{str(col).strip()}: {cell_str}")
                    except Exception:
                        continue
                if row_text_parts:
                    formatted_text = f"Sheet: {sheet_name}\n" + " | ".join(row_text_parts)
                    actual_row = df.index.get_loc(original_idx) + 1
                    text_data.append({'text': formatted_text, 'page_number': f"Row{actual_row}(Sheet:'{sheet_name}')", 'is_table': True})
                    all_text += formatted_text + "\n"
        workbook.close()
        if not text_data:
            return {}
        return {'filename': excel_path.name, 'title': excel_path.stem, 'text_data': text_data, 'full_text': all_text}
    except Exception as e:
        print(f"Error extracting from {excel_path}: {e}")
        return {}

def extract_file(file_path):
    ext = file_path.suffix.lower()
    if ext == '.pdf':
        return extract_content_from_pdf(file_path)
    elif ext in ('.docx', '.doc'):
        return extract_content_from_docx(file_path)
    elif ext == '.txt':
        return extract_content_from_txt(file_path)
    elif ext == '.json':
        return extract_content_from_json(file_path)
    elif ext in ('.xlsx', '.xls'):
        return extract_content_from_excel(file_path)
    else:
        print(f"Skipping unsupported: {ext}")
        return {}

# ============================================================
# CHUNKING
# ============================================================
def chunk_text(text, chunk_size=None, overlap=None):
    chunk_size = chunk_size or CHUNK_SIZE
    overlap = overlap or CHUNK_OVERLAP
    splitter = RecursiveCharacterTextSplitter(chunk_size=chunk_size, chunk_overlap=overlap)
    return splitter.split_text(text)

# ============================================================
# METADATA FUNCTIONS
# ============================================================
def extract_doi_from_pdf(pdf_path):
    try:
        with pymupdf.open(pdf_path) as pdf_doc:
            first_page = pdf_doc.load_page(0)
            text = first_page.get_text("text")
            doi_regex = r'\b(10.\d{4,9}/[-._;()/:A-Z0-9]+)\b'
            match = re.search(doi_regex, text, re.IGNORECASE)
            if match:
                return match.group(0)
    except Exception:
        pass
    return None

def fetch_crossref_metadata(doi):
    import requests
    url = f"https://api.crossref.org/works/{doi}"
    try:
        response = requests.get(url, timeout=10)
        if response.status_code == 200:
            data = response.json()["message"]
            authors = ", ".join([f"{a.get('given', '')} {a.get('family', '')}".strip() for a in data.get("author", [])])
            date_parts = data.get("published-print", data.get("published-online", {})).get("date-parts", [[]])[0]
            if isinstance(date_parts, list) and len(date_parts) >= 1:
                year = str(date_parts[0])
                month = str(date_parts[1]).zfill(2) if len(date_parts) >= 2 else '01'
                day = str(date_parts[2]).zfill(2) if len(date_parts) >= 3 else '01'
                published_date = f"{year}-{month}-{day}"
            else:
                published_date = "Not found"
            citation_count = data.get("is-referenced-by-count", 0)
            citation = f"{data.get('title', [''])[0]} by {authors} ({published_date}), DOI: {doi}"
            if citation_count > 0:
                citation += f" [Cited by: {citation_count}]"
            return {
                "doi": doi, "title": data.get("title", [""])[0], "authors": authors,
                "published": published_date, "citation_count": citation_count, "Citation": citation
            }
    except Exception:
        pass
    return {}

def extract_title_from_text(client, text):
    try:
        snippet = text[:4000]
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": "You are a strict title extractor. Return only the document's exact title. If you cannot identify one, return an empty string. No explanations."},
                {"role": "user", "content": f"TEXT:\n{snippet}"}
            ],
            temperature=0.0, max_tokens=64
        )
        title = (response.choices[0].message.content or "").strip().strip('`#\'"\"\"')
        title = title.splitlines()[0].strip() if title else ""
        return "" if len(title) > 400 else title
    except Exception as e:
        logging.error(f"Failed to extract title: {e}")
        return ""

def generate_summary(client, text_data):
    try:
        trimmed = text_data[:8000]
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": "Provide a concise summary of this document."},
                {"role": "user", "content": trimmed}
            ],
            max_tokens=300
        )
        return response.choices[0].message.content
    except Exception as e:
        print(f"Failed to generate summary: {e}")
        return "Summary not available."

# ============================================================
# EMBEDDING
# ============================================================
def generate_embeddings(client, text):
    try:
        response = client.embeddings.create(input=[text], model=EMBEDDING_MODEL)
        return response.data[0].embedding
    except Exception as e:
        print(f"Error generating embedding: {e}")
        return None

# ============================================================
# CSV GENERATION
# ============================================================
def create_csv(client):
    out_path = OUTPUT_FOLDER / CSV_FILENAME
    files = [f for f in INPUT_FOLDER.iterdir() if f.is_file()]
    print(f"Found {len(files)} files to process")

    columns = [
        "ID", "SOURCE_FILE", "CHUNK_INDEX", "CHUNK_PREVIEW", "TEXT", "PAGES",
        "CITATION_COUNT", "DOI", "TITLE", "AUTHORS", "PUBLISHED", "CITATION",
        "PAGE_REFERENCE", "EMBEDDING", "SAS_URL", "IS_TABLE", "FILE_TYPE", "SUMMARY"
    ]

    with open(out_path, "w", encoding="utf-8", newline="") as fp:
        writer = csv.DictWriter(fp, fieldnames=columns)
        writer.writeheader()

        for file_path in tqdm(files, desc="Processing files"):
            document = extract_file(file_path)
            if not document or not document.get('text_data'):
                continue

            # Get DOI and metadata for PDFs
            doi = None
            meta_crossref = {}
            if file_path.suffix.lower() == ".pdf":
                doi = extract_doi_from_pdf(file_path)
                if doi:
                    meta_crossref = fetch_crossref_metadata(doi)

            # Extract title using AI
            full_text = document.get('full_text', '') or "\n".join([d['text'] for d in document['text_data']])
            ai_title = extract_title_from_text(client, full_text)
            final_title = meta_crossref.get('title') or ai_title or document.get('title', 'Untitled')

            # Generate summary
            summary = generate_summary(client, full_text)

            # Process each text segment
            for page_data in document['text_data']:
                text = page_data.get('text', '').strip()
                if not text:
                    continue

                chunks = chunk_text(text)
                if not chunks:
                    continue

                for i, chunk_text_content in enumerate(chunks):
                    if not chunk_text_content.strip():
                        continue

                    chunk_id = f"{file_path.name}-{page_data.get('page_number', 1)}-{i}"
                    embedding = generate_embeddings(client, chunk_text_content)
                    if embedding is None:
                        continue

                    row = {
                        'ID': chunk_id,
                        'SOURCE_FILE': file_path.name,
                        'CHUNK_INDEX': i,
                        'CHUNK_PREVIEW': chunk_text_content[:200],
                        'TEXT': chunk_text_content,
                        'PAGES': json.dumps([str(page_data.get('page_number', 1))]),
                        'CITATION_COUNT': meta_crossref.get("citation_count", 0),
                        'DOI': meta_crossref.get("doi", doi or ""),
                        'TITLE': final_title,
                        'AUTHORS': meta_crossref.get("authors", ""),
                        'PUBLISHED': meta_crossref.get("published", ""),
                        'CITATION': meta_crossref.get("Citation", ""),
                        'PAGE_REFERENCE': f"p. {page_data.get('page_number', 1)}",
                        'EMBEDDING': json.dumps(embedding),
                        'SAS_URL': "",
                        'IS_TABLE': page_data.get('is_table', False),
                        'FILE_TYPE': file_path.suffix.replace('.', '').lower(),
                        'SUMMARY': summary
                    }
                    writer.writerow(row)

    print(f"CSV saved to {out_path}")
    return out_path

# ============================================================
# SNOWFLAKE UPLOAD
# ============================================================
def upload_csv_to_snowflake(csv_file_path):
    try:
        ctx = get_snowflake_connection()
        if ctx is None:
            return False

        cur = ctx.cursor()
        df = pd.read_csv(csv_file_path)
        print(f"Read CSV with {len(df)} rows")

        df.columns = df.columns.str.upper()

        # Clean PUBLISHED dates
        def clean_date(date_str):
            if pd.isna(date_str) or date_str in ['Not found', 'n/a', '', 'Not found']:
                return None
            try:
                pd.to_datetime(date_str)
                return date_str
            except Exception:
                return None

        df['PUBLISHED'] = df['PUBLISHED'].apply(clean_date)

        if RECREATE_TABLE:
            print(f"Recreating table {TABLE_NAME}...")
            cur.execute(f"DROP TABLE IF EXISTS {TABLE_NAME}")

        create_table_sql = f"""
        CREATE TABLE IF NOT EXISTS {TABLE_NAME} (
            ID VARCHAR(16777216),
            SOURCE_FILE VARCHAR(16777216),
            CHUNK_INDEX INTEGER,
            CHUNK_PREVIEW VARCHAR(16777216),
            TEXT VARCHAR(16777216),
            PAGES VARCHAR(16777216),
            CITATION_COUNT INTEGER,
            DOI VARCHAR(16777216),
            TITLE VARCHAR(16777216),
            AUTHORS VARCHAR(16777216),
            PUBLISHED DATE,
            CITATION VARCHAR(16777216),
            PAGE_REFERENCE VARCHAR(16777216),
            EMBEDDING VARCHAR(16777216),
            EMBEDDING_VECTOR VECTOR(FLOAT, 1536),
            SAS_URL VARCHAR(16777216),
            IS_TABLE BOOLEAN,
            FILE_TYPE VARCHAR(16777216),
            SUMMARY VARCHAR(16777216)
        )
        """
        cur.execute(create_table_sql)
        print(f"Table {TABLE_NAME} ready")

        if UPLOAD_MODE == "replace":
            cur.execute(f"TRUNCATE TABLE {TABLE_NAME}")
            print(f"Cleared {TABLE_NAME}")

        # Ensure warehouse is active before COPY INTO (write_pandas uses new internal cursors)
        warehouse = SNOW_CONFIG.get('warehouse') or 'WH_COMMUNICATIONS__EU__DER'
        cur.execute(f"USE WAREHOUSE {warehouse}")

        # Upload data
        success, nchunks, nrows, _ = write_pandas(
            conn=ctx, df=df, table_name=TABLE_NAME,
            database=SNOW_CONFIG['database'], schema=SNOW_CONFIG['schema'],
            chunk_size=1000, compression='gzip', on_error='continue', parallel=4
        )

        if success:
            print(f"Uploaded {nrows} rows to {TABLE_NAME}")
        else:
            print("Failed to upload data")
            return False

        # Convert embeddings to VECTOR format
        print("Converting embeddings to VECTOR format...")
        cur.execute(f"""
            UPDATE {TABLE_NAME}
            SET EMBEDDING_VECTOR = PARSE_JSON(EMBEDDING)::VECTOR(FLOAT, 1536)
            WHERE EMBEDDING IS NOT NULL AND EMBEDDING_VECTOR IS NULL
        """)
        print(f"Converted {cur.rowcount} embeddings to vectors")

        # Verify
        cur.execute(f"SELECT COUNT(*) FROM {TABLE_NAME}")
        total = cur.fetchone()[0]
        cur.execute(f"SELECT COUNT(*) FROM {TABLE_NAME} WHERE EMBEDDING_VECTOR IS NOT NULL")
        vectors = cur.fetchone()[0]
        print(f"Final: {total} total rows, {vectors} with vectors")

        return True
    except Exception as e:
        print(f"Error uploading to Snowflake: {e}")
        import traceback
        traceback.print_exc()
        return False
    finally:
        if 'cur' in locals():
            cur.close()
        if 'ctx' in locals():
            ctx.close()

# ============================================================
# MAIN
# ============================================================
if __name__ == "__main__":
    print("=" * 60)
    print("RAG ETL Pipeline - Local")
    print("=" * 60)
    print(f"Input:  {INPUT_FOLDER}")
    print(f"Output: {OUTPUT_FOLDER}")
    print(f"Table:  {TABLE_NAME}")
    print(f"Mode:   {UPLOAD_MODE}")
    print()

    # Test connections
    client = get_openai_client()
    if client is None:
        print("Cannot proceed without OpenAI")
        exit(1)

    if UPLOAD_TO_SNOWFLAKE:
        if not test_snowflake_connection():
            print("Cannot proceed without Snowflake")
            exit(1)

    print("\nAll connections OK\n")

    # Generate CSV
    if GENERATE_NEW_CSV:
        csv_path = create_csv(client)
    else:
        csv_path = OUTPUT_FOLDER / CSV_FILENAME

    # Upload to Snowflake
    if UPLOAD_TO_SNOWFLAKE:
        if csv_path.exists() if isinstance(csv_path, Path) else Path(csv_path).exists():
            success = upload_csv_to_snowflake(csv_path)
            if success:
                print("\nPipeline complete! Data uploaded to Snowflake.")
            else:
                print("\nUpload failed.")
        else:
            print(f"CSV not found at {csv_path}")

    print("\nDone!")
