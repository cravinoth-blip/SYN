"""
Patient Journey Pipeline — Main Orchestrator

This is the entry point. It runs the 4-pass pipeline sequentially,
managing the tool-calling loop for each pass, passing outputs forward,
and compiling the final deliverables + audit trail.

Usage:
    python orchestrator.py --disease "Systemic Lupus Erythematosus"
    python orchestrator.py --disease "Rheumatoid Arthritis" --supplements data/comp.xlsx data/epi.csv
"""

import os
import sys
import re
import json
import uuid
import argparse
import time
from datetime import datetime, timezone
from dataclasses import dataclass
from typing import Optional

from openai import OpenAI

import config
from audit import AuditLogger
from tools import build_tool_harness
from tools.base import ToolHarness
from passes.pass1_generate import get_pass1_prompt, get_pass1_user_message
from passes.pass2_verify import PASS2_SYSTEM_PROMPT, get_pass2_user_message
from passes.pass3_artifacts import PASS3_SYSTEM_PROMPT, get_pass3_user_message
from passes.pass4_polish import PASS4_SYSTEM_PROMPT, get_pass4_user_message


# --- Result Container -------------------------------------------------------

@dataclass
class PipelineResult:
    disease: str
    run_id: str
    pass1_json: dict
    pass2_json: dict
    pass3_json: dict
    polished_markdown: str
    deliverable_path: str       # Final .docx
    artifact_paths: list[str]   # Excel/CSV files from Pass 3
    audit_trail_path: str       # Audit .md
    audit_json_path: str        # Audit .json
    total_tool_calls: int
    duration_seconds: float


# --- Core Orchestrator -------------------------------------------------------

class PatientJourneyPipeline:

    def __init__(self):
        self.client = OpenAI(api_key=config.OPENAI_API_KEY)

    # -------------------------------------------------------------------------
    # Public entry point
    # -------------------------------------------------------------------------

    def run(
        self,
        disease: str,
        plan: dict = None,
        supplements: list[str] = None,
    ) -> PipelineResult:
        """
        Execute the full 4-pass pipeline for a given disease.

        Args:
            disease: Disease/condition name (e.g. "Systemic Lupus Erythematosus")
            plan: Optional locked plan dict (pre-approved research plan)
            supplements: Optional list of file paths to CI supplement Excel/CSV files

        Returns:
            PipelineResult with all outputs and file paths
        """
        # Fix #9: UTC-aware run_id and timestamps
        run_id = datetime.now(timezone.utc).strftime("%Y%m%d_%H%M%S") + "_" + uuid.uuid4().hex[:6]
        start_time = time.time()

        print(f"\n{'='*70}")
        print(f"  Patient Journey Pipeline - {disease}")
        print(f"  Run ID: {run_id}")
        print(f"{'='*70}\n")

        # Setup
        audit = AuditLogger(disease=disease, run_id=run_id)
        harness = build_tool_harness(audit=audit, supplements=supplements or [])

        os.makedirs(config.OUTPUT_DIR, exist_ok=True)
        os.makedirs(config.AUDIT_LOG_DIR, exist_ok=True)

        # Pass 1: Deep Generation
        print(">> PASS 1 - Deep Generation")
        pass1_json = self._run_pass(
            pass_number=1,
            system_prompt=get_pass1_prompt(disease),
            user_message=get_pass1_user_message(disease, plan),
            harness=harness,
            audit=audit,
            max_tool_calls=config.MAX_TOOL_CALLS_PASS1,
            model=config.MODEL_GENERATION,
        )
        # Fix #6: validate required keys
        self._require_keys(pass1_json, ["phases"], pass_number=1)
        print(f"  Pass 1 complete - {len(pass1_json.get('phases', []))} phases generated\n")

        # Pass 2: Verification & Deepening
        print(">> PASS 2 - Verification & Deepening")
        pass2_json = self._run_pass(
            pass_number=2,
            system_prompt=PASS2_SYSTEM_PROMPT,
            user_message=get_pass2_user_message(pass1_json),
            harness=harness,
            audit=audit,
            max_tool_calls=config.MAX_TOOL_CALLS_PASS2,
            model=config.MODEL_GENERATION,
        )
        # Fix #6: validate required keys
        self._require_keys(pass2_json, ["phases"], pass_number=2)
        print(f"  Pass 2 complete - confidence levels assigned\n")

        # Pass 3: Artifact Construction
        print(">> PASS 3 - Artifact Construction")
        pass3_json = self._run_pass(
            pass_number=3,
            system_prompt=PASS3_SYSTEM_PROMPT,
            user_message=get_pass3_user_message(pass2_json),
            harness=harness,
            audit=audit,
            max_tool_calls=config.MAX_TOOL_CALLS_PASS3,
            model=config.MODEL_ARTIFACTS,
        )
        # Fix #6: validate required keys
        self._require_keys(pass3_json, ["built_artifacts"], pass_number=3)
        artifact_paths = self._collect_artifact_paths(pass3_json)
        print(f"  Pass 3 complete - {len(artifact_paths)} artifacts built\n")

        # Pass 4: Editorial Polish
        print(">> PASS 4 - Editorial Polish")
        polished_md = self._run_pass_text(
            pass_number=4,
            system_prompt=PASS4_SYSTEM_PROMPT,
            user_message=get_pass4_user_message(pass2_json, pass3_json),
            audit=audit,
            model=config.MODEL_POLISH,
        )
        print(f"  Pass 4 complete - polished narrative generated\n")

        # Convert to Word Document
        deliverable_path = self._markdown_to_docx(polished_md, disease, run_id)

        # Save Audit Trail
        total_tool_calls = len(audit.entries)
        audit_md_path = audit.save()
        audit_json_path = audit.save_raw_json()

        duration = time.time() - start_time

        print(f"{'='*70}")
        print(f"  Pipeline complete in {duration:.1f}s")
        print(f"  Total tool calls: {total_tool_calls}")
        print(f"  Deliverable: {deliverable_path}")
        print(f"  Artifacts: {artifact_paths}")
        print(f"  Audit trail: {audit_md_path}")
        print(f"{'='*70}\n")

        return PipelineResult(
            disease=disease,
            run_id=run_id,
            pass1_json=pass1_json,
            pass2_json=pass2_json,
            pass3_json=pass3_json,
            polished_markdown=polished_md,
            deliverable_path=deliverable_path,
            artifact_paths=artifact_paths,
            audit_trail_path=audit_md_path,
            audit_json_path=audit_json_path,
            total_tool_calls=total_tool_calls,
            duration_seconds=duration,
        )

    # -------------------------------------------------------------------------
    # Pass runner (with tool loop)
    # -------------------------------------------------------------------------

    def _run_pass(
        self,
        pass_number: int,
        system_prompt: str,
        user_message: str,
        harness: ToolHarness,
        audit: AuditLogger,
        max_tool_calls: int,
        model: str,
    ) -> dict:
        """
        Run a single pass with the full tool-calling loop.

        The model can call tools iteratively. We keep looping until:
        - The model returns a final text response (no more tool calls), or
        - We hit the max_tool_calls safety cap

        Returns: parsed JSON from the model's final response.
        """
        messages = [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_message},
        ]

        tool_schemas = harness.get_openai_tools()
        tool_call_count = 0
        tools_used = []
        key_decisions = []

        while tool_call_count < max_tool_calls:
            # Fix #2: use retry helper instead of direct client call
            response = self._call_model(
                model=model,
                messages=messages,
                tools=tool_schemas,
                tool_choice="auto",
                temperature=config.TEMPERATURE,
                max_tokens=config.MAX_TOKENS_PER_PASS,
            )

            choice = response.choices[0]

            # Model is done (no tool calls)
            if choice.finish_reason == "stop" or not choice.message.tool_calls:
                final_text = choice.message.content or ""
                # Fix #10: record decision
                key_decisions.append("Model completed without further tool use")

                audit.log_pass_summary(
                    pass_number=pass_number,
                    total_tool_calls=tool_call_count,
                    tools_used=tools_used,
                    key_decisions=key_decisions,
                    output_summary=f"Completed with {tool_call_count} tool calls",
                )
                return self._extract_json(final_text)

            assistant_msg = choice.message
            messages.append(assistant_msg)

            for tool_call in assistant_msg.tool_calls:
                # Fix #1: enforce cap inside the inner loop
                if tool_call_count >= max_tool_calls:
                    print(f"    Tool budget reached at {tool_call_count}/{max_tool_calls}")
                    key_decisions.append(f"Tool budget reached at {tool_call_count}")
                    break

                tool_name = tool_call.function.name
                try:
                    arguments = json.loads(tool_call.function.arguments)
                except json.JSONDecodeError:
                    arguments = {}
                    # Fix #10: record bad JSON args
                    key_decisions.append(f"Invalid JSON arguments for {tool_name}")

                tools_used.append(tool_name)
                tool_call_count += 1

                print(f"    [{tool_call_count:2d}] {tool_name}({self._truncate_args(arguments)})")

                # Fix #3: safe tool dispatch
                try:
                    result = harness.dispatch(
                        tool_name=tool_name,
                        arguments=arguments,
                        pass_number=pass_number,
                    )
                except Exception as e:
                    result = {
                        "error": True,
                        "tool_name": tool_name,
                        "arguments": arguments,
                        "message": f"{type(e).__name__}: {str(e)}",
                    }
                    # Fix #10: record tool failure
                    key_decisions.append(f"Tool failure: {tool_name}")

                # Fix #5: structured truncation instead of blind slice
                serialized = json.dumps(result, default=str)
                if len(serialized) > 20000:
                    serialized = json.dumps({
                        "truncated": True,
                        "message": "Tool output exceeded 20000 chars",
                        "preview": serialized[:20000],
                    })

                messages.append({
                    "role": "tool",
                    "tool_call_id": tool_call.id,
                    "content": serialized,
                })

            # Fix #1: break outer loop once cap is hit
            if tool_call_count >= max_tool_calls:
                break

        # Safety cap reached — force a final completion call
        print(f"    Max tool calls ({max_tool_calls}) reached, forcing completion")
        key_decisions.append(f"Forced completion after hitting cap: {max_tool_calls}")

        # Fix #2: use retry helper
        response = self._call_model(
            model=model,
            messages=messages + [
                {"role": "user", "content": "Tool call budget exhausted. Return your best output now as JSON."}
            ],
            temperature=config.TEMPERATURE,
            max_tokens=config.MAX_TOKENS_PER_PASS,
        )

        audit.log_pass_summary(
            pass_number=pass_number,
            total_tool_calls=tool_call_count,
            tools_used=tools_used,
            key_decisions=key_decisions,
            output_summary=f"Completed (hit cap at {max_tool_calls} calls)",
        )

        return self._extract_json(response.choices[0].message.content or "")

    # -------------------------------------------------------------------------
    # Pass 4 runner (text-only, no tools)
    # -------------------------------------------------------------------------

    def _run_pass_text(
        self,
        pass_number: int,
        system_prompt: str,
        user_message: str,
        audit: AuditLogger,
        model: str,
    ) -> str:
        """Run Pass 4 - pure text transformation, no tools needed."""
        # Fix #2: use retry helper
        response = self._call_model(
            model=model,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_message},
            ],
            temperature=0.4,
            max_tokens=config.MAX_TOKENS_PER_PASS,
        )

        text = response.choices[0].message.content or ""

        audit.log_pass_summary(
            pass_number=4,
            total_tool_calls=0,
            tools_used=[],
            key_decisions=["Pure text transformation pass"],
            output_summary=f"Generated {len(text)} chars of polished markdown",
        )

        return text

    # -------------------------------------------------------------------------
    # Fix #2: model call retry helper
    # -------------------------------------------------------------------------

    def _call_model(self, **kwargs):
        """Call the OpenAI chat completions API with up to 3 retries on failure."""
        last_error = None
        for attempt in range(3):
            try:
                return self.client.chat.completions.create(**kwargs)
            except Exception as e:
                last_error = e
                wait = 2 ** attempt
                print(f"    Model call failed (attempt {attempt + 1}/3): {e}")
                time.sleep(wait)
        raise RuntimeError(f"Model call failed after 3 attempts: {last_error}")

    # -------------------------------------------------------------------------
    # Fix #4: robust JSON extraction
    # -------------------------------------------------------------------------

    def _extract_json(self, text: str) -> dict:
        """
        Extract JSON object or array from model response.
        Returns a dict; arrays are wrapped for consistency.
        """
        if not text or not text.strip():
            return {"raw_response": "", "parse_error": True, "reason": "empty_response"}

        text = text.strip()

        # Prefer fenced json blocks
        for marker in ["```json", "```"]:
            if marker in text:
                try:
                    start = text.index(marker) + len(marker)
                    end = text.index("```", start)
                    candidate = text[start:end].strip()
                    parsed = json.loads(candidate)
                    return parsed if isinstance(parsed, dict) else {"data": parsed}
                except Exception:
                    pass

        # Try the full text directly
        try:
            parsed = json.loads(text)
            return parsed if isinstance(parsed, dict) else {"data": parsed}
        except Exception:
            pass

        # Scan for the first valid JSON object or array using the streaming decoder
        decoder = json.JSONDecoder()
        for i, ch in enumerate(text):
            if ch not in "{[":
                continue
            try:
                parsed, _ = decoder.raw_decode(text[i:])
                return parsed if isinstance(parsed, dict) else {"data": parsed}
            except Exception:
                continue

        print(f"    Could not parse JSON from model response ({len(text)} chars)")
        return {"raw_response": text, "parse_error": True}

    # -------------------------------------------------------------------------
    # Fix #6: pass output validator
    # -------------------------------------------------------------------------

    def _require_keys(self, payload: dict, required: list[str], pass_number: int) -> None:
        """Raise if any required key is missing from a pass output."""
        missing = [k for k in required if k not in payload]
        if missing:
            raise ValueError(
                f"Pass {pass_number} output missing required keys: {missing}. "
                f"Payload keys: {list(payload.keys())}"
            )

    # -------------------------------------------------------------------------
    # Markdown -> Word conversion
    # -------------------------------------------------------------------------

    def _markdown_to_docx(self, markdown: str, disease: str, run_id: str) -> str:
        """Convert polished markdown to a formatted Word document."""
        try:
            from docx import Document
            from docx.shared import Inches, Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            path = os.path.join(config.OUTPUT_DIR, f"journey_{run_id}.md")
            with open(path, "w", encoding="utf-8") as f:
                f.write(markdown)
            print(f"    python-docx not installed, saved as markdown: {path}")
            return path

        doc = Document()

        style = doc.styles["Normal"]
        font = style.font
        font.name = "Calibri"
        font.size = Pt(11)

        title = doc.add_heading("Patient Journey Map", level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle.add_run(disease)
        run.font.size = Pt(18)
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

        meta = doc.add_paragraph()
        meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Fix #9: UTC timestamp
        meta.add_run(
            f"Generated: {datetime.now(timezone.utc).strftime('%B %d, %Y %H:%M UTC')}"
        ).font.size = Pt(10)

        doc.add_page_break()

        for line in markdown.split("\n"):
            line = line.strip()
            if not line:
                continue
            elif line.startswith("### "):
                doc.add_heading(line[4:], level=3)
            elif line.startswith("## "):
                doc.add_heading(line[3:], level=2)
            elif line.startswith("# "):
                doc.add_heading(line[2:], level=1)
            elif line.startswith("- "):
                doc.add_paragraph(line[2:], style="List Bullet")
            # Fix #7: regex-based numbered list (handles any number, not just 1/2/3)
            elif re.match(r"^\d+\.\s+", line):
                text = re.sub(r"^\d+\.\s+", "", line)
                doc.add_paragraph(text, style="List Number")
            else:
                doc.add_paragraph(line)

        path = os.path.join(config.OUTPUT_DIR, f"patient_journey_{run_id}.docx")
        doc.save(path)
        return path

    # -------------------------------------------------------------------------
    # Artifact collection
    # -------------------------------------------------------------------------

    def _collect_artifact_paths(self, pass3_output: dict) -> list[str]:
        """Extract file paths of artifacts created by Pass 3."""
        paths = []
        for artifact in pass3_output.get("built_artifacts", []):
            fpath = artifact.get("file_path", "")
            if fpath and os.path.isfile(fpath):
                paths.append(fpath)
            elif fpath:
                workspace_path = os.path.join("./workspace", os.path.basename(fpath))
                if os.path.isfile(workspace_path):
                    paths.append(workspace_path)
        return paths

    # -------------------------------------------------------------------------
    # Utilities
    # -------------------------------------------------------------------------

    @staticmethod
    def _truncate_args(args: dict, max_len: int = 80) -> str:
        """Truncate tool arguments for console display."""
        s = ", ".join(f"{k}={repr(v)[:30]}" for k, v in args.items())
        return s[:max_len] + "..." if len(s) > max_len else s


# --- CLI Entry Point ---------------------------------------------------------

def main():
    parser = argparse.ArgumentParser(description="Run the Patient Journey Pipeline")
    parser.add_argument(
        "--disease", "-d",
        required=True,
        help="Disease or condition name (e.g. 'Systemic Lupus Erythematosus')",
    )
    parser.add_argument(
        "--supplements", "-s",
        nargs="*",
        default=[],
        help="Paths to CI supplement files (Excel/CSV)",
    )
    parser.add_argument(
        "--plan", "-p",
        default=None,
        help="Path to a locked plan JSON file (optional)",
    )

    args = parser.parse_args()

    # Fix #8: validate file inputs early
    plan = None
    if args.plan:
        if not os.path.isfile(args.plan):
            raise FileNotFoundError(f"Plan file not found: {args.plan}")
        with open(args.plan, "r", encoding="utf-8") as f:
            plan = json.load(f)

    for s in args.supplements:
        if not os.path.isfile(s):
            raise FileNotFoundError(f"Supplement file not found: {s}")

    pipeline = PatientJourneyPipeline()
    result = pipeline.run(
        disease=args.disease,
        plan=plan,
        supplements=args.supplements,
    )

    print(f"\nDeliverable:  {result.deliverable_path}")
    print(f"Artifacts:    {result.artifact_paths}")
    print(f"Audit trail:  {result.audit_trail_path}")
    print(f"Tool calls:   {result.total_tool_calls}")
    print(f"Duration:     {result.duration_seconds:.1f}s\n")


if __name__ == "__main__":
    main()
